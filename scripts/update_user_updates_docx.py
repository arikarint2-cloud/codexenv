from __future__ import annotations

from datetime import date
from pathlib import Path

from docx import Document

DOCX_PATH = Path('docs/user-updates.docx')

UPDATE = {
    'title': f"Обновления от {date.today().strftime('%d.%m.%Y')}",
    'sections': {
        'Что нового': [
            'Главный экран стал информативнее: расширен первый блок с понятным описанием пользы консультаций и ключевых преимуществ.'
        ],
        'Улучшения': [
            'Навигация на сайте стала удобнее на телефоне: меню лучше реагирует на действия и проще закрывается.',
            'Улучшена читабельность текстов и элементов интерфейса на мобильных устройствах.',
            'Страница контактов получила более понятные формулировки, чтобы быстрее сориентироваться по следующим шагам.'
        ],
        'Исправления': [
            'Исправлены проблемы с отображением русскоязычного текста в отдельных частях сайта.',
            'Убрано вводящее в заблуждение поведение формы: теперь явно указано, что это демонстрационный режим.'
        ],
        'Важно знать': [
            'Отправка формы пока не выполняет реальную передачу данных в рабочую систему.'
        ]
    }
}


def add_update_block(doc: Document, payload: dict) -> int:
    doc.add_heading(payload['title'], level=1)
    count = 0

    for category, items in payload['sections'].items():
        if not items:
            continue
        doc.add_heading(category, level=2)
        for item in items:
            doc.add_paragraph(item, style='List Bullet')
            count += 1

    doc.add_paragraph('')
    return count


def main() -> None:
    DOCX_PATH.parent.mkdir(parents=True, exist_ok=True)

    if DOCX_PATH.exists():
        doc = Document(DOCX_PATH)
    else:
        doc = Document()

    added = add_update_block(doc, UPDATE)
    doc.save(DOCX_PATH)
    print(f'updated={DOCX_PATH} items={added}')


if __name__ == '__main__':
    main()
